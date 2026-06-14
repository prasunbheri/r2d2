from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('R2 Control Bridge', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Raspberry Pi Zero W + 2 × Cytron MDD20A — 4-Motor Control with Camera & Virtual Joystick', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# 1. Overview
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'This project turns a Raspberry Pi Zero W into a WiFi-enabled motor controller for four '
    'brushed DC motors. Two Cytron MDD20A dual-channel H-bridge drivers provide 20A per channel '
    '(60A peak) for high-power motors. A web-based dashboard served from the Pi lets a phone or '
    'laptop control all four motors in real time over WiFi, with sub-15ms latency. '
    'An OV5647 camera module provides a live MJPEG video feed in the dashboard.'
)

p = doc.add_paragraph('Key capabilities:')
p.runs[0].bold = True
caps = [
    'Virtual joystick with differential drive mapping (FL/RL = y+x, FR/RR = y-x)',
    'Auto-Center toggle + adjustable return-to-center damping (Joystick Speed)',
    'Max Speed Limiter scales motor output proportionally (0–100%)',
    'Live OV5647 camera feed (640×480 MJPEG) embedded in dashboard',
    'Real-time control via WebSocket (~10–15ms round trip)',
    'Smooth PWM at 20kHz (inaudible, no motor whine)',
    'Emergency stop (instant kill all motors)',
    'Reconnection overlay with countdown when WebSocket drops',
    'Auto-start on boot — headless operation, no monitor or keyboard needed',
]
for c in caps:
    doc.add_paragraph(c, style='List Bullet')

# 2. Hardware
doc.add_heading('2. Hardware Components', level=1)

doc.add_heading('2.1 Bill of Materials', level=2)
table = doc.add_table(rows=9, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Item', 'Component', 'Qty', 'Notes']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['1', 'Raspberry Pi Zero W', '1', 'WiFi, 512MB RAM, runs Flask + pigpio'],
    ['2', 'Cytron MDD20A', '2', '20A/chan dual H-bridge, 6-30V supply'],
    ['3', 'DC motors (brushed)', '4', 'Up to 20A continuous each'],
    ['4', 'Battery / PSU (6-30V)', '1', 'Powers the motors via MDD20A'],
    ['5', 'Jumper wires (female-female)', '10', '8 signal + 2 ground for GPIO'],
    ['6', 'MicroSD card (8GB+)', '1', 'Raspbian OS installed'],
    ['7', 'LED (any color) + 220Ω resistor', '1', 'Status indicator on GPIO5'],
    ['8', 'OV5647 camera module', '1', 'CSI interface, 640×480 MJPEG stream'],
]
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')
doc.add_heading('2.2 Wiring Diagram (GPIO Pin Assignment)', level=2)

doc.add_paragraph('All connections use the bottom 10 pins of the Pi Zero W 40-pin header:')

table2 = doc.add_table(rows=11, cols=4, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['Physical Pin', 'BCM GPIO', 'Signal', 'Connected To']
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

wiring = [
    ['31', 'GPIO6',  'DIR — Front Left',     'MDD20A #1 — DIR1'],
    ['32', 'GPIO12', 'PWM — Front Left',      'MDD20A #1 — PWM1'],
    ['33', 'GPIO13', 'DIR — Front Right',     'MDD20A #1 — DIR2'],
    ['34', 'GND',    'Ground',                'MDD20A #1 — GND'],
    ['35', 'GPIO19', 'PWM — Front Right',     'MDD20A #1 — PWM2'],
    ['36', 'GPIO16', 'DIR — Rear Left',       'MDD20A #2 — DIR1'],
    ['37', 'GPIO26', 'PWM — Rear Left',       'MDD20A #2 — PWM1'],
    ['38', 'GPIO20', 'DIR — Rear Right',      'MDD20A #2 — DIR2'],
    ['39', 'GND',    'Ground',                'MDD20A #2 — GND'],
    ['40', 'GPIO21', 'PWM — Rear Right',      'MDD20A #2 — PWM2'],
]
for row_idx, row_data in enumerate(wiring, start=1):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')
doc.add_paragraph('Important wiring rules:', style='List Bullet')
p = doc.add_paragraph(
    'Share GND between both drivers and the Pi — this is mandatory for signal reference.',
    style='List Bullet'
)
p = doc.add_paragraph(
    'Connect motor battery (6-30V) to each MDD20A\'s VB+/VB- terminals. '
    'The MDD20A has NO reverse-polarity protection — double-check polarity before connecting.',
    style='List Bullet'
)
p = doc.add_paragraph(
    'Power the Pi Zero W separately via its microUSB port (5V, 2.5A recommended).',
    style='List Bullet'
)
p = doc.add_paragraph(
    'All selected GPIO pins have no active alternate function conflicts on a standard Pi Zero W configuration.',
    style='List Bullet'
)

doc.add_heading('2.3 Status LED', level=2)
doc.add_paragraph(
    'An LED (with a 220Ω series resistor) is connected from GPIO5 to GND. '
    'This LED indicates the system state through blink patterns controlled by the watchdog service.'
)

led_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
led_table.alignment = WD_TABLE_ALIGNMENT.CENTER
led_data = [
    ['LED Pattern', 'System State'],
    ['Off', 'No power / Pi booting'],
    ['Slow blink (1s on, 1s off)', 'Stack starting — services not all active yet'],
    ['Solid on', 'All services running — system ready'],
    ['Fast blink (200ms on, 200ms off)', 'Service crash detected — watchdog is restarting it'],
    ['3 quick flashes then pause', 'Fatal error — watchdog cannot recover the service'],
]
for i, (col1, col2) in enumerate(led_data):
    led_table.rows[i].cells[0].text = col1
    led_table.rows[i].cells[1].text = col2
    if i == 0:
        for p in led_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in led_table.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.bold = True

# 3. Software Architecture
doc.add_heading('3. Software Architecture', level=1)

doc.add_heading('3.1 Technology Stack', level=2)
stack_table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stack_data = [
    ['Layer', 'Technology'],
    ['PWM Generation', 'pigpio — DMA-based, hardware-timed PWM on any GPIO (20kHz)'],
    ['Web Server', 'Flask + Flask-SocketIO — lightweight, async WebSocket support'],
    ['Camera', 'picamera2 — OV5647 sensor, 640×480 MJPEG stream on /video_feed'],
    ['Frontend', 'Plain HTML/CSS/JS — virtual joystick, touch-friendly, no framework dependencies'],
    ['Socket.IO Client', 'v4.7.5 served locally from static/socket.io.min.js (no CDN dependency)'],
    ['System Service', 'systemd — auto-starts pigpiod + web server + watchdog on boot'],
]
for i, (col1, col2) in enumerate(stack_data):
    stack_table.rows[i].cells[0].text = col1
    stack_table.rows[i].cells[1].text = col2
    if i == 0:
        for p in stack_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in stack_table.rows[i].cells[1].paragraphs:
            for r in p.runs:
                r.bold = True

doc.add_paragraph('')
doc.add_heading('3.2 Files on the Pi', level=2)
file_table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
file_table.alignment = WD_TABLE_ALIGNMENT.CENTER
file_headers = ['File', 'Purpose', 'Key Details']
for i, h in enumerate(file_headers):
    file_table.rows[0].cells[i].text = h
    for p in file_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

file_data = [
    ['motor_control.py', 'Hardware abstraction layer', 'Connects to pigpiod, set_speed(motor, -100..100), stop_all(), 20kHz DMA PWM'],
    ['app.py', 'Flask + SocketIO web server', 'Binds 0.0.0.0:5000, camera init, MJPEG /video_feed, WebSocket events'],
    ['templates/index.html', 'Web dashboard UI', 'Virtual joystick, Auto-Center, Speed/Max Speed sliders, camera, reconnect overlay'],
    ['static/socket.io.min.js', 'SocketIO client library', 'v4.7.5 — served locally so CDN is not required'],
    ['watchdog.py', 'Service monitor + LED controller', 'Polls service states every 2s, controls GPIO5 LED patterns, restarts failed services'],
    ['motor_control.service', 'systemd unit for web server', 'After=pigpiod.service, Restart=on-failure, RestartSec=5'],
    ['watchdog.service', 'systemd unit for watchdog', 'After=motor_control.service pigpiod.service, Requires=pigpiod.service'],
    ['deploy.sh', 'Deployment script', 'Copies all files and installs systemd services'],
]
for row_idx, row_data in enumerate(file_data, start=1):
    for col_idx, val in enumerate(row_data):
        file_table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')
doc.add_heading('3.3 Control Flow', level=2)
doc.add_paragraph(
    '1. User drags the virtual joystick in the phone browser\n'
    '2. JavaScript computes target speeds via differential drive: FL/RL = y+x, FR/RR = y-x\n'
    '3. rAF animation loop interpolates motor speeds toward target (return-to-center damping)\n'
    '4. Throttled to 30 updates/sec, sent via WebSocket as set_speeds event\n'
    '5. Flask-SocketIO receives the event with 4 motor values\n'
    '6. app.py calls motor_control.set_speed(motor_id, value) for each motor\n'
    '7. motor_control.py converts speed to DIR (high/low) + PWM duty cycle (0-1000)\n'
    '8. pigpio writes to GPIO registers via /dev/gpiomem (DMA, sub-1ms)\n'
    '9. MDD20A drives each motor at the commanded speed\n\n'
    'Total latency: ~10-15ms round trip — perceived as instant by the user.'
)

# 4. Features
doc.add_heading('4. Features', level=1)

features = [
    ('Virtual Joystick with Differential Drive',
     'A single virtual joystick (touch circle) controls all 4 motors via differential drive mapping: '
     'FL/RL = y+x, FR/RR = y-x. Vertical tilt drives forward/backward, horizontal tilt turns. '
     'The joystick renders as a draggable dot within a circular boundary, showing real-time X/Y values.'),
    ('Auto-Center with Adjustable Damping',
     'When enabled, the joystick returns to center on release and motor speeds smoothly glide back to zero. '
     'The Joystick Speed slider (50–3050ms) controls the return-to-center damping time — higher values '
     'give a slower, more gradual deceleration. When disabled, motors hold their last commanded speed.'),
    ('Max Speed Limiter',
     'Scales the motor output proportionally from 0% to 100%. At 50% (default), the joystick\'s full throw '
     'produces only 50% PWM duty cycle. Useful for limiting speed during testing or indoor use without '
     'changing the joystick sensitivity curve.'),
    ('Live Camera Feed',
     'An OV5647 camera module (CSI interface) streams 640×480 MJPEG video via a dedicated /video_feed '
     'endpoint. The feed auto-reconnects when WebSocket reconnection occurs (cache-busting timestamp). '
     'The MJPEG generator uses stream_with_context + GeneratorExit handling to prevent thread exhaustion '
     'on disconnect/reconnect loops. Positioned below the connection status in the dashboard.'),
    ('Emergency Stop',
     'A prominent STOP button kills all motor outputs immediately by setting all PWM duty cycles to zero. '
     'This is a software emergency stop — for physical safety, consider adding a hardware kill switch '
     'inline with the motor battery.'),
    ('Reconnection Overlay',
     'When the WebSocket disconnects, the dashboard dims all controls to 40% opacity and shows a glowing '
     'red "Reconnecting" overlay with a countdown timer ("in N seconds"). Retries every 5 seconds. '
     'On reconnection, the overlay disappears and full control is restored.'),
    ('20kHz Silent PWM',
     'The PWM frequency is set to 20kHz, above the range of human hearing. Motors run silently without '
     'the whining noise typical of lower-frequency PWM (e.g., 1kHz from RPi.GPIO software PWM).'),
    ('Local Socket.IO Serving',
     'The Socket.IO client library (v4.7.5) is served from the Pi itself at static/socket.io.min.js. '
     'No CDN dependency — works completely offline, no internet connection required on the phone.'),
    ('Headless Auto-Start',
     'On power-up, the Pi automatically starts the pigpio daemon, Flask web server, and watchdog service. '
     'No monitor, keyboard, or manual login required — just power on and connect.'),
    ('IP Announcement on Startup',
     'The web server prints the Pi\'s IP address to the console on startup. When accessed via serial or SSH, '
     'the user immediately sees which IP to connect to.'),
    ('Touch-Friendly Interface',
     'The dashboard is designed for phone use: large virtual joystick (200px), responsive layout that '
     'scrolls vertically on small screens, thick touch targets, camera feed capped at 240px width.'),
    ('Modular Codebase',
     'motor_control.py is a standalone class that can be imported and used from other Python scripts, '
     'not just the web server. This allows future expansion (e.g., autonomous mode, ROS integration).'),
]

for t, d in features:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

# 5. Use Cases
doc.add_heading('5. Use Cases', level=1)

usecases = [
    ('Remote-Controlled Robot Vehicle',
     'Control a 4-wheel or tracked robot from your phone. The differential-drive joystick enables '
     'intuitive driving: forward/backward/turn-in-place. The Pi Zero W and camera mount on the robot chassis.'),
    ('Multi-Motor Test Bench',
     'Test and characterize multiple motors individually or simultaneously via the independent motor API.'),
    ('Wireless Conveyor / Material Handling',
     'Control a 4-motor conveyor system from a safe distance with individual speed control.'),
    ('Educational Platform',
     'Learn motor control, PWM, web services, computer vision, and embedded Linux all in one project.'),
    ('Prototype Development',
     'Rapid prototyping platform for motor control algorithms. WebSocket interface makes it easy to write '
     'custom control UIs or integrate with higher-level systems.'),
]

for t, d in usecases:
    doc.add_heading(t, level=2)
    doc.add_paragraph(d)

# 6. Installation
doc.add_heading('6. Installation & Setup', level=1)

doc.add_heading('6.1 Prerequisites', level=2)
doc.add_paragraph('Raspbian OS (or Raspberry Pi OS) installed and running on the Pi Zero W.')
doc.add_paragraph('Pi connected to the internet (for package installation).')
doc.add_paragraph('SSH access to the Pi (username/password provided).')
doc.add_paragraph('OV5647 camera module connected via CSI ribbon cable.')

doc.add_heading('6.2 Step-by-Step', level=2)
steps = [
    'Install system packages: sudo apt update && sudo apt install -y pigpio python3-pip python3-picamera2',
    'Enable pigpio daemon: sudo systemctl enable pigpiod && sudo systemctl start pigpiod',
    'Install Python packages: pip3 install flask flask-socketio',
    'Verify camera: libcamera-hello --timeout 2000 (should show camera preview)',
    'Wire the Pi to the MDD20A drivers per the wiring table in Section 2.2',
    'Connect the status LED (with 220Ω resistor) from GPIO5 to GND',
    'Connect motors to the MDD20A output terminals',
    'Connect motor battery (6-30V) to MDD20A VB+/VB-',
    'Power the Pi Zero W via microUSB',
    'Copy all source files to /home/r2tele/motor_control/:\n'
    '  motor_control.py, app.py, watchdog.py, deploy.sh\n'
    '  templates/index.html\n'
    '  static/socket.io.min.js',
    'Create and enable systemd services:\n'
    '  sudo cp motor_control.service /etc/systemd/system/\n'
    '  sudo cp watchdog.service /etc/systemd/system/\n'
    '  sudo systemctl daemon-reload\n'
    '  sudo systemctl enable motor_control.service watchdog.service',
    'Reboot: sudo reboot',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('6.3 Connecting from a Phone', level=2)
doc.add_paragraph('With the Pi connected to your phone\'s hotspot:')
doc.add_paragraph('Open a browser on your phone (Chrome recommended)')
doc.add_paragraph('Navigate to http://<pi-ip>:5000')
doc.add_paragraph('Use an IP scanner app (e.g., Fing) to find the Pi\'s IP if unknown')

doc.add_paragraph('With the Pi creating its own WiFi access point:')
doc.add_paragraph('Connect your phone to the Pi\'s WiFi network')
doc.add_paragraph('Open http://192.168.4.1:5000 in the browser')

# 7. Safety
doc.add_heading('7. Safety Notes', level=1)
safety_items = [
    'The MDD20A has NO reverse-polarity protection on the motor power input. Connecting the battery backwards will instantly destroy the driver.',
    'High-current motors can draw >20A — use appropriately thick wires (minimum 12 AWG recommended for the battery-to-driver connection).',
    'The software emergency stop stops PWM signals, but the motors may coast. For mission-critical safety, add a physical relay or contactor that cuts motor battery power.',
    'The Pi Zero W GPIO pins are 3.3V and not 5V tolerant. The MDD20A accepts 1.8-12V logic, so this is compatible. Do not connect 5V logic outputs to the Pi GPIOs.',
    'Heat: at 20A continuous, the MDD20A may generate significant heat. Ensure adequate ventilation. The MDD20A uses NMOS H-bridges and typically doesn\'t require heatsinks at rated currents.',
]
for item in safety_items:
    doc.add_paragraph(item, style='List Bullet')

# 8. Watchdog & LED System
doc.add_heading('8. Watchdog & LED Status System', level=1)

doc.add_heading('8.1 Overview', level=2)
doc.add_paragraph(
    'The watchdog service runs continuously in the background to monitor the health of all critical '
    'services (pigpiod, motor_control web server). It uses the GPIO5 LED to visually indicate the '
    'system state at a glance. If a service crashes, the watchdog automatically restarts it and '
    'signals the fault through the LED.'
)

doc.add_heading('8.2 Service Dependency Chain', level=2)
doc.add_paragraph(
    'Power-on → pigpiod starts (systemd)\n'
    '        → motor_control (app.py) starts after pigpiod\n'
    '        → watchdog starts after pigpiod (NOT bound to motor_control)\n'
    '        → watchdog confirms all alive → GPIO5 solid on'
)

doc.add_paragraph(
    'If any service crashes:\n'
    '  • systemd Restart=on-failure attempts restart (primary mechanism)\n'
    '  • watchdog detects the failure within 2 seconds (secondary/backup)\n'
    '  • watchdog runs systemctl restart if systemd did not recover it\n'
    '  • LED switches to fast blink during recovery\n'
    '  • After 3 failed restart attempts, LED shows fatal error pattern\n\n'
    'Note: watchdog.service does NOT include motor_control.service in its Requires= directive. '
    'This was intentional — stopping motor_control manually (e.g., for maintenance) will not kill the watchdog. '
    'The watchdog can then detect the stoppage and restart motor_control automatically.'
)

doc.add_heading('8.3 Watchdog Logic', level=2)
doc.add_paragraph(
    '1. On start: GPIO5 slow blink — system booting\n'
    '2. Poll all required services every 2 seconds\n'
    '3. All active → GPIO5 solid on (ready)\n'
    '4. Any service down → GPIO5 fast blink + systemctl restart\n'
    '5. Re-check in 5 seconds\n'
    '6. Restored → back to solid on\n'
    '7. Still down after 3 retries → 3-flash fatal pattern\n'
    '8. Watchdog itself restarts automatically via systemd if it crashes'
)

doc.add_heading('8.4 Edge Cases Handled', level=2)
edge_cases = [
    'Watchdog crashes → systemd restarts it (Restart=on-failure in watchdog.service)',
    'pigpiod frozen (not crashed) → systemd watchdog timeout kills and restarts it',
    'Race condition on boot → motor_control waits for pigpiod socket before connecting',
    'Network down but motors still work → watchdog only checks services, not network',
    'Manual service stop → watchdog survives (not coupled) and auto-restarts motor_control',
    'MJPEG thread exhaustion → GeneratorExit handler cleans up threads on disconnect',
]
for item in edge_cases:
    doc.add_paragraph(item, style='List Bullet')

# 9. Future Enhancements
doc.add_heading('9. Future Enhancements', level=1)
future = [
    'Encoder feedback for closed-loop PID speed control',
    'Battery voltage monitoring via MDD20A or external ADC',
    'WiFi AP mode: Pi creates its own network so phone connects directly without a hotspot',
    'MQTT integration for IoT / home automation connectivity',
    'ROS 2 integration for advanced robotics applications',
    'OLED display for IP address and status (no SSH required)',
]
for item in future:
    doc.add_paragraph(item, style='List Bullet')

# 10. Project Structure & Test Suite
doc.add_heading('10. Project Structure & Test Suite', level=1)

doc.add_heading('10.1 File Tree', level=2)
tree = (
    'motor_control/\n'
    '├── motor_control.py          # Hardware abstraction — pigpio DMA PWM, 4 motors\n'
    '├── app.py                    # Flask + SocketIO + camera MJPEG server\n'
    '├── watchdog.py               # Service monitor + GPIO5 LED controller\n'
    '├── motor_control.service     # systemd unit for web server\n'
    '├── watchdog.service          # systemd unit for watchdog\n'
    '├── deploy.sh                 # One-shot deploy script\n'
    '├── templates/\n'
    '│   └── index.html            # Virtual joystick UI, camera feed, reconnect overlay\n'
    '├── static/\n'
    '│   └── socket.io.min.js      # SocketIO client v4.7.5 (local, no CDN)\n'
    '└── tests/\n'
    '    ├── __init__.py\n'
    '    ├── mock_pigpio.py        # Mock pigpio library for testing\n'
    '    ├── test_motor_control.py  # 29 tests — init, speed, direction, clamping, cleanup\n'
    '    ├── test_app.py            # 7 tests — HTTP routes, WebSocket events\n'
    '    └── test_watchdog.py       # 14 tests — LED patterns, service detection, restart\n'
)
doc.add_paragraph(tree.replace('  ', '\u00a0\u00a0'))

doc.add_heading('10.2 Test Coverage Summary', level=2)
test_table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
test_headers = ['Test File', 'Tests', 'What It Covers']
for i, h in enumerate(test_headers):
    test_table.rows[0].cells[i].text = h
    for p in test_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
test_data = [
    ['test_motor_control.py', '29', 'Init, set_speed (±100), clamping, set_all, stop, cleanup, error handling, independent channels'],
    ['test_app.py', '7', 'HTTP 200, WebSocket set_speed/set_speeds/stop, invalid motor ignored, connect sends state'],
    ['test_watchdog.py', '14', 'LED on/off/blink patterns, service_active detection, restart call, wait_for_services timeout'],
]
for row_idx, row_data in enumerate(test_data, start=1):
    for col_idx, val in enumerate(row_data):
        test_table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph('')
doc.add_paragraph(
    'All 50 tests pass without hardware — they use mock_pigpio.py which simulates the pigpio '
    'library. Tests can be run on any machine (no Pi required) with:\n'
    '  cd motor_control && pip install pytest && python3 -m pytest tests/ -v'
)

# Save
doc.save('/home/prasun/AI/git/r2/R2_Motor_Control_Plan.docx')
print("Document updated: R2_Motor_Control_Plan.docx")
